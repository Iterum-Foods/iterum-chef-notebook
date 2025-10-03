"""
Enhanced Menu Parser Service
Improved parsing for Word documents and various menu formats
"""

import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
import json
from docx import Document
import io

logger = logging.getLogger(__name__)

@dataclass
class EnhancedMenuItem:
    """Enhanced menu item with better parsing"""
    title: str
    description: str = ""
    price: Optional[float] = None
    price_text: str = ""
    category: str = ""
    dietary_tags: List[str] = None
    spice_level: Optional[str] = None
    allergens: List[str] = None
    ingredients: List[str] = None
    confidence: float = 0.0
    
    def __post_init__(self):
        if self.dietary_tags is None:
            self.dietary_tags = []
        if self.allergens is None:
            self.allergens = []
        if self.ingredients is None:
            self.ingredients = []

@dataclass
class EnhancedMenuSection:
    """Enhanced menu section with better organization"""
    name: str
    items: List[EnhancedMenuItem]
    description: str = ""
    confidence: float = 0.0

class EnhancedMenuParser:
    """Enhanced menu parser with better Word document support"""
    
    def __init__(self):
        # Enhanced price patterns
        self.price_patterns = [
            r'\$(\d+(?:\.\d{2})?)',  # $12.50, $12
            r'(\d+(?:\.\d{2})?)\s*(?:dollars?|USD|\$)',  # 12.50 dollars, 12 USD
            r'(\d+(?:\.\d{2})?)\s*(?:£|pounds?)',  # 12.50 pounds, £12
            r'(\d+(?:\.\d{2})?)\s*(?:€|euros?)',  # 12.50 euros, €12
            r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:won|₩)',  # Korean won
            r'(\d+(?:\.\d{2})?)\s*(?:yen|¥)',  # Japanese yen
            r'(\d+(?:\.\d{2})?)\s*(?:CAD|AUD|NZD)',  # Other currencies
            r'(\d+(?:\.\d{2})?)\s*(?:per\s+person|pp)',  # Per person pricing
            r'(\d+(?:\.\d{2})?)\s*(?:each|ea)',  # Each pricing
        ]
        
        # Enhanced section patterns
        self.section_patterns = [
            # Common section headers
            r'^(APPETIZERS?|STARTERS?|SMALL\s+PLATES?|HORS\s+D\'OEUVRE)',
            r'^(SOUPS?|SALADS?|GREEN\s+SALADS?)',
            r'^(MAIN\s+COURSES?|ENTREES?|MAINS?|ENTRÉES?)',
            r'^(DESSERTS?|SWEETS?|DESSERT\s+MENU)',
            r'^(BEVERAGES?|DRINKS?|COCKTAILS?|WINE\s+LIST)',
            r'^(SIDES?|SIDE\s+DISHES?|ACCOMPANIMENTS?)',
            r'^(SPECIALS?|CHEF\'?S?\s+SPECIALS?|DAILY\s+SPECIALS?)',
            r'^(LUNCH|DINNER|BREAKFAST|BRUNCH|HAPPY\s+HOUR)',
            r'^(KIDS?\s+MENU|CHILDREN\'?S?\s+MENU)',
            r'^(VEGETARIAN|VEGAN|PLANT\s+BASED)',
            r'^(GLUTEN\s+FREE|DAIRY\s+FREE)',
            r'^(SEAFOOD|FISH|MEAT|POULTRY)',
            r'^(PASTA|PIZZA|BURGERS?|SANDWICHES?)',
            # Generic patterns
            r'^([A-Z][A-Z\s&]+)(?:\s*---|\s*–|\s*—|\s*\.\.\.|\s*:).*$',
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*$',  # Title case headers
        ]
        
        # Enhanced dietary patterns
        self.dietary_patterns = {
            'vegetarian': r'\b(?:vegetarian|veggie|V\b|VEG)',
            'vegan': r'\b(?:vegan|VG\b|VGN)',
            'gluten_free': r'\b(?:gluten[\s-]?free|GF\b|GLUTEN\s+FREE)',
            'dairy_free': r'\b(?:dairy[\s-]?free|DF\b|DAIRY\s+FREE)',
            'nut_free': r'\b(?:nut[\s-]?free|NF\b|NUT\s+FREE)',
            'sugar_free': r'\b(?:sugar[\s-]?free|SF\b|SUGAR\s+FREE)',
            'keto': r'\b(?:keto|ketogenic|KETO)',
            'paleo': r'\b(?:paleo|paleolithic|PALEO)',
            'low_carb': r'\b(?:low[\s-]?carb|LOW\s+CARB)',
            'organic': r'\b(?:organic|ORG)',
            'local': r'\b(?:local|LOC)',
            'fresh': r'\b(?:fresh|FRESH)',
            'homemade': r'\b(?:homemade|house[\s-]?made|HOMEMADE)',
            'signature': r'\b(?:signature|chef\'?s?\s+special|house\s+special|SIG)',
            'popular': r'\b(?:popular|bestseller|customer\s+favorite|POP)',
            'new': r'\b(?:new|NEW)',
            'seasonal': r'\b(?:seasonal|SEASONAL)',
        }
        
        # Allergen patterns
        self.allergen_patterns = {
            'nuts': r'\b(?:nuts?|almonds?|walnuts?|pecans?|pistachios?|cashews?|hazelnuts?)',
            'dairy': r'\b(?:milk|cheese|butter|cream|yogurt|dairy)',
            'eggs': r'\b(?:eggs?|egg)',
            'soy': r'\b(?:soy|soybean|tofu)',
            'wheat': r'\b(?:wheat|flour|bread|pasta)',
            'fish': r'\b(?:fish|salmon|tuna|cod)',
            'shellfish': r'\b(?:shellfish|shrimp|crab|lobster|scallops?)',
            'sesame': r'\b(?:sesame|tahini)',
        }
        
        # Spice level patterns
        self.spice_patterns = {
            'mild': r'\b(?:mild|gentle|light\s+spice|1\s*chili|🌶)',
            'medium': r'\b(?:medium|moderate\s+spice|2\s*chili|🌶🌶)',
            'hot': r'\b(?:hot|spicy|very\s+spicy|3\s*chili|🌶🌶🌶)',
            'extra_hot': r'\b(?:extra\s+hot|extremely\s+spicy|4\s*chili|🌶🌶🌶🌶)',
            'fiery': r'\b(?:fiery|blazing|5\s*chili|🌶🌶🌶🌶🌶)'
        }
        
        # Menu item title patterns
        self.title_patterns = [
            r'^([A-Z][A-Z\s&]+)$',  # All caps titles
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)$',  # Title case
            r'^(\d+\.?\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',  # Numbered items
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*[-–—]',  # Title with dash
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*\.{2,}',  # Title with dots
        ]
    
    def parse_menu_text(self, text: str, source_type: str = "text") -> Dict[str, Any]:
        """
        Parse menu text with enhanced accuracy
        
        Args:
            text: Raw menu text
            source_type: Type of source (text, word, pdf, image)
            
        Returns:
            Dictionary with parsed menu structure
        """
        try:
            # Clean and normalize text based on source type
            cleaned_text = self.clean_text(text, source_type)
            
            # Extract sections with enhanced logic
            sections = self.extract_sections_enhanced(cleaned_text)
            
            # Parse each section with improved accuracy
            parsed_sections = []
            total_items = 0
            
            for section_name, section_text in sections:
                menu_items = self.extract_menu_items_enhanced(section_text, section_name)
                if menu_items:
                    confidence = self.calculate_section_confidence(menu_items)
                    parsed_sections.append(EnhancedMenuSection(
                        name=section_name,
                        items=menu_items,
                        confidence=confidence
                    ))
                    total_items += len(menu_items)
            
            # If no clear sections found, try alternative parsing
            if not parsed_sections:
                all_items = self.extract_menu_items_enhanced(cleaned_text, "Menu Items")
                if all_items:
                    confidence = self.calculate_section_confidence(all_items)
                    parsed_sections.append(EnhancedMenuSection(
                        name="Menu Items",
                        items=all_items,
                        confidence=confidence
                    ))
                    total_items = len(all_items)
            
            # Create enhanced summary
            summary = self.create_enhanced_summary(parsed_sections)
            
            return {
                'success': True,
                'sections': [self.section_to_dict(section) for section in parsed_sections],
                'summary': summary,
                'total_items': total_items,
                'total_sections': len(parsed_sections),
                'parsing_info': {
                    'method': f'enhanced_menu_parsing_{source_type}',
                    'confidence': self.calculate_overall_confidence(parsed_sections, text),
                    'status': 'completed',
                    'source_type': source_type
                }
            }
            
        except Exception as e:
            logger.error(f"Enhanced menu parsing failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'sections': [],
                'summary': {},
                'total_items': 0,
                'parsing_info': {'status': 'failed', 'method': f'enhanced_menu_parsing_{source_type}'}
            }
    
    def parse_word_document(self, doc_content: bytes) -> Dict[str, Any]:
        """
        Parse Word document with enhanced table and formatting support
        
        Args:
            doc_content: Word document content as bytes
            
        Returns:
            Dictionary with parsed menu structure
        """
        try:
            # Load Word document
            doc = Document(io.BytesIO(doc_content))
            
            # Extract text with formatting information
            text_content = []
            table_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract table content
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    table_content.append(table_data)
            
            # Combine text and table content
            combined_text = text_content.copy()
            
            # Process tables as structured data
            for table in table_content:
                combined_text.append("")  # Add separator
                for row in table:
                    combined_text.append(" | ".join(row))
            
            full_text = "\n".join(combined_text)
            
            # Parse with enhanced parser
            return self.parse_menu_text(full_text, "word")
            
        except Exception as e:
            logger.error(f"Word document parsing failed: {e}")
            return {
                'success': False,
                'error': f"Word document parsing failed: {str(e)}",
                'sections': [],
                'summary': {},
                'total_items': 0,
                'parsing_info': {'status': 'failed', 'method': 'word_document_parsing'}
            }
    
    def clean_text(self, text: str, source_type: str = "text") -> str:
        """Enhanced text cleaning based on source type"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Pipe to I
        text = text.replace('0', 'O')  # Zero to O in some contexts
        text = text.replace('l', 'I')  # Lowercase l to uppercase I
        
        # Fix Word document specific issues
        if source_type == "word":
            # Remove Word-specific characters
            text = text.replace('\u2013', '-')  # En dash
            text = text.replace('\u2014', '--')  # Em dash
            text = text.replace('\u2018', "'")  # Left single quote
            text = text.replace('\u2019', "'")  # Right single quote
            text = text.replace('\u201c', '"')  # Left double quote
            text = text.replace('\u201d', '"')  # Right double quote
            text = text.replace('\u2026', '...')  # Ellipsis
        
        # Clean up table separators
        text = re.sub(r'\|\s*\|\s*\|', ' | ', text)  # Multiple pipes
        text = re.sub(r'\s*\|\s*', ' | ', text)  # Normalize pipe spacing
        
        return text.strip()
    
    def extract_sections_enhanced(self, text: str) -> List[Tuple[str, str]]:
        """Enhanced section extraction with better pattern matching"""
        sections = []
        lines = text.split('\n')
        current_section = "Menu Items"
        current_content = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            is_section_header = False
            section_name = None
            
            for pattern in self.section_patterns:
                match = re.match(pattern, line.upper())
                if match:
                    # Save previous section
                    if current_content:
                        sections.append((current_section, '\n'.join(current_content)))
                    
                    # Start new section
                    section_name = self.normalize_section_name(line)
                    current_section = section_name
                    current_content = []
                    is_section_header = True
                    break
            
            # Additional heuristics for section detection
            if not is_section_header:
                # Check if line looks like a section header
                if self.looks_like_section_header(line, lines, i):
                    if current_content:
                        sections.append((current_section, '\n'.join(current_content)))
                    
                    current_section = self.normalize_section_name(line)
                    current_content = []
                    is_section_header = True
            
            if not is_section_header:
                current_content.append(line)
        
        # Add final section
        if current_content:
            sections.append((current_section, '\n'.join(current_content)))
        
        return sections
    
    def looks_like_section_header(self, line: str, all_lines: List[str], line_index: int) -> bool:
        """Enhanced section header detection"""
        # All caps and short
        if line.isupper() and len(line) < 50 and len(line.split()) <= 4:
            return True
        
        # Title case and isolated
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', line) and len(line.split()) <= 4:
            # Check if next line is not a menu item
            if line_index + 1 < len(all_lines):
                next_line = all_lines[line_index + 1].strip()
                if not self.looks_like_menu_item(next_line):
                    return True
        
        # Contains common section keywords
        section_keywords = ['menu', 'appetizer', 'main', 'dessert', 'drink', 'special', 'lunch', 'dinner']
        if any(keyword in line.lower() for keyword in section_keywords):
            return True
        
        return False
    
    def extract_menu_items_enhanced(self, text: str, section_name: str = "") -> List[EnhancedMenuItem]:
        """Enhanced menu item extraction with better accuracy"""
        items = []
        lines = text.split('\n')
        current_item = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check if this looks like a menu item
            if self.looks_like_menu_item(line):
                # Save previous item
                if current_item:
                    items.append(current_item)
                
                # Start new item
                current_item = self.create_menu_item_from_line(line, section_name)
                
            elif current_item and self.looks_like_description(line):
                # This line is likely a description continuation
                if current_item.description:
                    current_item.description += " " + line
                else:
                    current_item.description = line
                
                # Update dietary tags and allergens from description
                current_item.dietary_tags.extend(self.extract_dietary_tags(line))
                current_item.allergens.extend(self.extract_allergens(line))
        
        # Add final item
        if current_item:
            items.append(current_item)
        
        return items
    
    def looks_like_menu_item(self, line: str) -> bool:
        """Enhanced menu item detection"""
        # Has price
        if self.extract_price(line)['has_price']:
            return True
        
        # Matches title patterns
        for pattern in self.title_patterns:
            if re.match(pattern, line):
                return True
        
        # All caps and reasonable length
        if line.isupper() and 3 <= len(line) <= 50:
            return True
        
        # Title case and reasonable length
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*', line) and len(line.split()) <= 8:
            return True
        
        # Numbered items
        if re.match(r'^\d+\.?\s+[A-Z]', line):
            return True
        
        return False
    
    def looks_like_description(self, line: str) -> bool:
        """Check if line looks like a description"""
        # Not a title pattern
        if any(re.match(pattern, line) for pattern in self.title_patterns):
            return False
        
        # Not a section header
        if any(re.match(pattern, line.upper()) for pattern in self.section_patterns):
            return False
        
        # Not a price-only line
        if re.match(r'^\$?\d+(?:\.\d{2})?$', line):
            return False
        
        # Contains descriptive words
        descriptive_words = [
            'with', 'and', 'served', 'topped', 'drizzled', 'garnished', 'fresh', 'seasoned',
            'accompanied', 'includes', 'features', 'made', 'prepared', 'cooked', 'grilled',
            'roasted', 'baked', 'steamed', 'sautéed', 'braised', 'marinated', 'glazed',
            'sauce', 'dressing', 'garnish', 'side', 'topped with', 'served with', 'comes with'
        ]
        if any(word in line.lower() for word in descriptive_words):
            return True
        
        # Longer than typical titles
        if len(line) > 20:
            return True
        
        return False
    
    def create_menu_item_from_line(self, line: str, section_name: str = "") -> EnhancedMenuItem:
        """Create menu item from a line with enhanced parsing"""
        # Extract price
        price_info = self.extract_price(line)
        
        # Extract title
        title = self.extract_title_enhanced(line)
        
        # Extract description
        description = self.extract_description_enhanced(line, title, price_info['text'])
        
        # Extract dietary tags and allergens
        dietary_tags = self.extract_dietary_tags(line)
        allergens = self.extract_allergens(line)
        
        # Extract spice level
        spice_level = self.extract_spice_level(line)
        
        # Extract ingredients
        ingredients = self.extract_ingredients(line)
        
        # Calculate confidence
        confidence = self.calculate_item_confidence(line, price_info, title, description)
        
        return EnhancedMenuItem(
            title=title,
            description=description,
            price=price_info['price'],
            price_text=price_info['text'],
            category=section_name,
            dietary_tags=dietary_tags,
            spice_level=spice_level,
            allergens=allergens,
            ingredients=ingredients,
            confidence=confidence
        )
    
    def extract_title_enhanced(self, line: str) -> str:
        """Enhanced title extraction"""
        # Remove price first
        price_info = self.extract_price(line)
        line_without_price = line.replace(price_info['text'], '').strip()
        
        # Remove dietary indicators
        for pattern in self.dietary_patterns.values():
            line_without_price = re.sub(pattern, '', line_without_price, flags=re.IGNORECASE)
        
        # Remove allergen indicators
        for pattern in self.allergen_patterns.values():
            line_without_price = re.sub(pattern, '', line_without_price, flags=re.IGNORECASE)
        
        # Extract title using patterns
        for pattern in self.title_patterns:
            match = re.match(pattern, line_without_price)
            if match:
                return match.group(1).strip()
        
        # Fallback: take first part before common separators
        separators = [' - ', ' – ', ' — ', '...', ' | ', ' • ']
        for sep in separators:
            if sep in line_without_price:
                return line_without_price.split(sep)[0].strip()
        
        # Fallback: take first N words
        words = line_without_price.split()
        if len(words) > 8:
            return ' '.join(words[:8]).strip()
        
        return line_without_price.strip()
    
    def extract_description_enhanced(self, line: str, title: str, price_text: str) -> str:
        """Enhanced description extraction"""
        # Remove title and price
        description = line.replace(title, '').replace(price_text, '').strip()
        
        # Remove dietary and allergen tags
        for pattern in self.dietary_patterns.values():
            description = re.sub(pattern, '', description, flags=re.IGNORECASE)
        
        for pattern in self.allergen_patterns.values():
            description = re.sub(pattern, '', description, flags=re.IGNORECASE)
        
        # Clean up separators - take text after the separator
        separators = [' - ', ' – ', ' — ', '...', ' | ', ' • ', ':', ';', '•']
        for sep in separators:
            if sep in description:
                parts = description.split(sep, 1)
                if len(parts) > 1:
                    description = parts[1].strip()
                break
        
        # Clean up common prefixes that might be left over
        prefixes_to_remove = ['with ', 'includes ', 'features ', 'served ', 'comes ']
        for prefix in prefixes_to_remove:
            if description.lower().startswith(prefix):
                description = description[len(prefix):].strip()
                break
        
        return description.strip()
    
    def extract_ingredients(self, text: str) -> List[str]:
        """Extract ingredients from text"""
        ingredients = []
        
        # Common ingredient patterns
        ingredient_patterns = [
            r'\b(?:chicken|beef|pork|fish|salmon|tuna|cod|shrimp|crab|lobster)\b',
            r'\b(?:tomato|onion|garlic|pepper|mushroom|spinach|lettuce|carrot|potato)\b',
            r'\b(?:cheese|milk|cream|butter|yogurt|sour\s+cream)\b',
            r'\b(?:rice|pasta|noodle|bread|flour|sugar|salt|pepper)\b',
            r'\b(?:olive\s+oil|vegetable\s+oil|sesame\s+oil)\b',
            r'\b(?:basil|oregano|thyme|rosemary|parsley|cilantro)\b',
        ]
        
        for pattern in ingredient_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            ingredients.extend(matches)
        
        return list(set(ingredients))  # Remove duplicates
    
    def extract_allergens(self, text: str) -> List[str]:
        """Extract allergen information"""
        allergens = []
        for allergen, pattern in self.allergen_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                allergens.append(allergen)
        return allergens
    
    def extract_price(self, text: str) -> Dict[str, Any]:
        """Enhanced price extraction"""
        for pattern in self.price_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    price_str = match.group(1) if match.groups() else match.group(0).replace('$', '')
                    price_value = float(price_str.replace(',', ''))
                    return {
                        'has_price': True,
                        'price': price_value,
                        'text': match.group(0)
                    }
                except ValueError:
                    continue
        
        return {'has_price': False, 'price': None, 'text': ''}
    
    def extract_dietary_tags(self, text: str) -> List[str]:
        """Extract dietary restriction tags"""
        tags = []
        for tag_name, pattern in self.dietary_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                tags.append(tag_name)
        return tags
    
    def extract_spice_level(self, text: str) -> Optional[str]:
        """Extract spice level information"""
        for level, pattern in self.spice_patterns.items():
            if re.search(pattern, text, re.IGNORECASE):
                return level
        return None
    
    def calculate_item_confidence(self, line: str, price_info: Dict, title: str, description: str) -> float:
        """Calculate confidence score for a menu item"""
        confidence = 0.0
        
        # Base confidence
        confidence += 20.0
        
        # Price found
        if price_info['has_price']:
            confidence += 30.0
        
        # Good title
        if title and len(title) > 2:
            confidence += 25.0
        
        # Description present
        if description and len(description) > 5:
            confidence += 15.0
        
        # Dietary tags
        if any(re.search(pattern, line, re.IGNORECASE) for pattern in self.dietary_patterns.values()):
            confidence += 10.0
        
        return min(confidence, 100.0)
    
    def calculate_section_confidence(self, items: List[EnhancedMenuItem]) -> float:
        """Calculate confidence score for a section"""
        if not items:
            return 0.0
        
        total_confidence = sum(item.confidence for item in items)
        return total_confidence / len(items)
    
    def calculate_overall_confidence(self, sections: List[EnhancedMenuSection], original_text: str) -> float:
        """Calculate overall parsing confidence"""
        if not sections:
            return 0.0
        
        total_items = sum(len(section.items) for section in sections)
        if total_items == 0:
            return 0.0
        
        # Average item confidence
        all_items = [item for section in sections for item in section.items]
        avg_item_confidence = sum(item.confidence for item in all_items) / len(all_items)
        
        # Section structure bonus
        structure_bonus = min(len(sections) * 5, 20)
        
        # Price coverage bonus
        items_with_prices = sum(1 for item in all_items if item.price)
        price_bonus = (items_with_prices / len(all_items)) * 15
        
        return min(avg_item_confidence + structure_bonus + price_bonus, 100.0)
    
    def normalize_section_name(self, section_name: str) -> str:
        """Normalize section names"""
        name = section_name.strip().title()
        name = re.sub(r'[^\w\s]', '', name)  # Remove special characters
        return name
    
    def create_enhanced_summary(self, sections: List[EnhancedMenuSection]) -> Dict[str, Any]:
        """Create enhanced summary of the parsed menu"""
        total_items = sum(len(section.items) for section in sections)
        items_with_prices = sum(1 for section in sections for item in section.items if item.price)
        
        # Price range
        all_prices = [item.price for section in sections for item in section.items if item.price]
        price_range = {
            'min': min(all_prices) if all_prices else None,
            'max': max(all_prices) if all_prices else None,
            'average': sum(all_prices) / len(all_prices) if all_prices else None
        }
        
        # Dietary analysis
        all_dietary_tags = []
        all_allergens = []
        for section in sections:
            for item in section.items:
                all_dietary_tags.extend(item.dietary_tags)
                all_allergens.extend(item.allergens)
        
        dietary_counts = {}
        for tag in all_dietary_tags:
            dietary_counts[tag] = dietary_counts.get(tag, 0) + 1
        
        allergen_counts = {}
        for allergen in all_allergens:
            allergen_counts[allergen] = allergen_counts.get(allergen, 0) + 1
        
        # Confidence analysis
        avg_confidence = sum(section.confidence for section in sections) / len(sections) if sections else 0
        
        return {
            'total_sections': len(sections),
            'total_items': total_items,
            'items_with_prices': items_with_prices,
            'price_coverage': (items_with_prices / total_items * 100) if total_items > 0 else 0,
            'price_range': price_range,
            'section_names': [section.name for section in sections],
            'dietary_analysis': {
                'dietary_tags': dict(sorted(dietary_counts.items(), key=lambda x: x[1], reverse=True)),
                'allergens': dict(sorted(allergen_counts.items(), key=lambda x: x[1], reverse=True))
            },
            'average_items_per_section': total_items / len(sections) if sections else 0,
            'average_confidence': avg_confidence,
            'parsing_quality': self.assess_parsing_quality(sections)
        }
    
    def assess_parsing_quality(self, sections: List[EnhancedMenuSection]) -> str:
        """Assess the overall quality of parsing"""
        if not sections:
            return "Poor"
        
        total_items = sum(len(section.items) for section in sections)
        items_with_prices = sum(1 for section in sections for item in section.items if item.price)
        avg_confidence = sum(section.confidence for section in sections) / len(sections)
        
        if avg_confidence >= 80 and items_with_prices / total_items >= 0.7:
            return "Excellent"
        elif avg_confidence >= 60 and items_with_prices / total_items >= 0.5:
            return "Good"
        elif avg_confidence >= 40:
            return "Fair"
        else:
            return "Poor"
    
    def section_to_dict(self, section: EnhancedMenuSection) -> Dict[str, Any]:
        """Convert EnhancedMenuSection to dictionary"""
        return {
            'name': section.name,
            'description': section.description,
            'items': [self.item_to_dict(item) for item in section.items],
            'item_count': len(section.items),
            'confidence': section.confidence
        }
    
    def item_to_dict(self, item: EnhancedMenuItem) -> Dict[str, Any]:
        """Convert EnhancedMenuItem to dictionary"""
        return {
            'title': item.title,
            'description': item.description,
            'price': item.price,
            'price_text': item.price_text,
            'category': item.category,
            'dietary_tags': item.dietary_tags,
            'spice_level': item.spice_level,
            'allergens': item.allergens,
            'ingredients': item.ingredients,
            'confidence': item.confidence
        }

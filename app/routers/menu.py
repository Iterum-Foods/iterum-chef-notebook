from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from fastapi.responses import JSONResponse
import os
from pathlib import Path
from PyPDF2 import PdfReader
import logging
import io
from docx import Document
from sqlalchemy.orm import Session
from datetime import datetime
import uuid

from app.database import get_db
from app.core.auth import get_current_user
from app.services.menu_parser import MenuParser
from app.services.enhanced_menu_parser import EnhancedMenuParser
try:
    from app.services.ocr_processor import OCRProcessor
except ImportError:
    from app.services.ocr_processor_fallback import OCRProcessorFallback as OCRProcessor

UPLOAD_DIR = Path("uploads/menus")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

router = APIRouter()

@router.get("/health")
def menu_health():
    return {"status": "ok", "message": "Menu router is alive"}

@router.post("/test-upload")
async def test_upload(file: UploadFile = File(...)):
    # Just echo back the filename and size for debugging
    try:
        content = await file.read()
        return {"filename": file.filename, "size": len(content)}
    except Exception as e:
        logging.exception("Test upload failed")
        raise HTTPException(status_code=500, detail=f"Test upload failed: {e}")

@router.post("/upload-document")
async def upload_menu_document(file: UploadFile = File(...)):
    """Upload and extract text from PDF or Word documents"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    filename_lower = file.filename.lower()
    if not (filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
        raise HTTPException(
            status_code=400, 
            detail="Only PDF and Word (.docx) files are supported. Legacy .doc files are not supported."
        )
    
    try:
        # Save the uploaded file
        save_path = UPLOAD_DIR / file.filename
        content = await file.read()
        
        try:
            with open(save_path, "wb") as f:
                f.write(content)
        except Exception as e:
            logging.exception("Failed to save uploaded file")
            raise HTTPException(status_code=500, detail=f"Failed to save file: {e}")
        
        # Extract text based on file type
        extracted_text = ""
        
        if filename_lower.endswith(".pdf"):
            try:
                reader = PdfReader(str(save_path))
                extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as e:
                logging.exception("Failed to extract text from PDF")
                raise HTTPException(status_code=500, detail=f"Failed to extract text from PDF: {e}")
                
        elif filename_lower.endswith(".docx"):
            try:
                doc = Document(io.BytesIO(content))
                extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            except Exception as e:
                logging.exception("Failed to extract text from Word document")
                raise HTTPException(status_code=500, detail=f"Failed to extract text from Word document: {e}")
        
        return JSONResponse({
            "filename": file.filename,
            "file_type": "pdf" if filename_lower.endswith(".pdf") else "docx",
            "text": extracted_text.strip() or "No text found in document. (Scanned image? Try OCR.)",
            "success": True
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logging.exception("Unexpected error in document upload")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {e}")

@router.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    """Legacy PDF upload endpoint - use /upload-document for new implementations"""
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed.")
    try:
        # Save the uploaded file
        save_path = UPLOAD_DIR / file.filename
        try:
            with open(save_path, "wb") as f:
                content = await file.read()
                f.write(content)
        except Exception as e:
            logging.exception("Failed to save uploaded PDF")
            raise HTTPException(status_code=500, detail=f"Failed to save PDF: {e}")
        # Extract text from PDF
        try:
            reader = PdfReader(str(save_path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            logging.exception("Failed to extract text from PDF")
            raise HTTPException(status_code=500, detail=f"Failed to extract text: {e}")
        return JSONResponse({
            "filename": file.filename,
            "text": text.strip() or "No text found in PDF. (Scanned image? Try OCR.)"
        })
    except HTTPException:
        raise
    except Exception as e:
        logging.exception("Unexpected error in upload_pdf")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {e}")


@router.post("/extract-and-parse")
async def extract_and_parse_menu(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user)
):
    """
    Two-stage menu processing:
    1. Extract text from uploaded menu (PDF/image)
    2. Parse into structured menu items with titles, descriptions, prices
    """
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    try:
        # Stage 1: Extract text from uploaded file
        logging.info(f"Stage 1: Extracting text from {file.filename}")
        
        # Save uploaded file temporarily
        temp_filename = f"{uuid.uuid4()}_{file.filename}"
        temp_path = UPLOAD_DIR / temp_filename
        
        content = await file.read()
        with open(temp_path, "wb") as f:
            f.write(content)
        
        # Determine processing method based on file type
        filename_lower = file.filename.lower()
        extracted_text = ""
        extraction_method = ""
        extraction_info = {}
        
        if filename_lower.endswith(('.jpg', '.jpeg', '.png', '.gif')):
            # Use OCR for images
            ocr_processor = OCRProcessor()
            ocr_result = ocr_processor.extract_text_from_file(str(temp_path), f"image/{filename_lower.split('.')[-1]}")
            
            if ocr_result['success']:
                extracted_text = ocr_result['text']
                extraction_method = "OCR"
                extraction_info = ocr_result['processing_info']
            else:
                raise HTTPException(status_code=500, detail=f"OCR failed: {ocr_result.get('error', 'Unknown error')}")
                
        elif filename_lower.endswith('.pdf'):
            # Try direct text extraction first, then OCR if needed
            try:
                reader = PdfReader(str(temp_path))
                extracted_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                extraction_method = "PDF Text Extraction"
                extraction_info = {"pages": len(reader.pages), "method": "direct"}
                
                # If no text found, try OCR
                if not extracted_text.strip():
                    ocr_processor = OCRProcessor()
                    ocr_result = ocr_processor.extract_text_from_file(str(temp_path), "application/pdf")
                    
                    if ocr_result['success']:
                        extracted_text = ocr_result['text']
                        extraction_method = "PDF OCR"
                        extraction_info = ocr_result['processing_info']
                    
            except Exception as e:
                # Fallback to OCR
                ocr_processor = OCRProcessor()
                ocr_result = ocr_processor.extract_text_from_file(str(temp_path), "application/pdf")
                
                if ocr_result['success']:
                    extracted_text = ocr_result['text']
                    extraction_method = "PDF OCR Fallback"
                    extraction_info = ocr_result['processing_info']
                else:
                    raise HTTPException(status_code=500, detail=f"PDF processing failed: {str(e)}")
                    
        elif filename_lower.endswith('.docx'):
            # Use enhanced Word document parser
            enhanced_parser = EnhancedMenuParser()
            word_result = enhanced_parser.parse_word_document(content)
            
            if word_result['success']:
                # Return the enhanced parsing result directly
                return {
                    "success": True,
                    "stage": 2,
                    "extraction_method": "Enhanced Word Document Parser",
                    "extraction_info": {"enhanced_parsing": True},
                    "sections": word_result['sections'],
                    "summary": word_result['summary'],
                    "total_items": word_result['total_items'],
                    "parsing_info": word_result['parsing_info']
                }
            else:
                # Fallback to basic Word extraction
                doc = Document(io.BytesIO(content))
                extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                extraction_method = "Word Document (Basic)"
                extraction_info = {"paragraphs": len(doc.paragraphs), "enhanced_failed": True}
            
        else:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Supported: PDF, DOCX, JPG, PNG"
            )
        
        # Clean up temporary file
        try:
            os.unlink(temp_path)
        except:
            pass
        
        if not extracted_text.strip():
            return {
                "success": False,
                "stage": 1,
                "error": "No text could be extracted from the file",
                "extraction_method": extraction_method,
                "suggested_action": "Try a higher quality scan or different file format"
            }
        
        # Stage 2: Parse extracted text into structured menu items
        logging.info(f"Stage 2: Parsing menu structure from extracted text")
        
        # Use enhanced parser for better accuracy
        enhanced_parser = EnhancedMenuParser()
        parsed_result = enhanced_parser.parse_menu_text(extracted_text, extraction_method.lower())
        
        if not parsed_result['success']:
            return {
                "success": False,
                "stage": 2,
                "error": parsed_result.get('error', 'Menu parsing failed'),
                "extracted_text": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
                "extraction_method": extraction_method
            }
        
        # Combine results
        return {
            "success": True,
            "filename": file.filename,
            "processing_stages": {
                "stage_1_extraction": {
                    "method": extraction_method,
                    "text_length": len(extracted_text),
                    "info": extraction_info
                },
                "stage_2_parsing": {
                    "method": "Menu Text Parsing",
                    "confidence": parsed_result['parsing_info']['confidence'],
                    "items_found": parsed_result['total_items']
                }
            },
            "menu_data": {
                "sections": parsed_result['sections'],
                "summary": parsed_result['summary'],
                "total_items": parsed_result['total_items'],
                "total_sections": parsed_result['total_sections']
            },
            "extracted_text": extracted_text,
            "parsing_confidence": parsed_result['parsing_info']['confidence']
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logging.exception("Unexpected error in menu extraction and parsing")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


@router.post("/parse-text")
async def parse_menu_text(
    menu_data: dict,
    current_user = Depends(get_current_user)
):
    """
    Parse already extracted menu text into structured items
    Input: {"text": "menu text content"}
    """
    
    try:
        text = menu_data.get("text", "")
        if not text:
            raise HTTPException(status_code=400, detail="No text provided")
        
        menu_parser = MenuParser()
        result = menu_parser.parse_menu_text(text)
        
        return {
            "success": result['success'],
            "menu_data": {
                "sections": result['sections'],
                "summary": result['summary'],
                "total_items": result['total_items'],
                "total_sections": result['total_sections']
            },
            "parsing_info": result['parsing_info']
        }
        
    except Exception as e:
        logging.exception("Menu text parsing failed")
        raise HTTPException(status_code=500, detail=f"Parsing failed: {str(e)}")


@router.get("/sample")
async def get_sample_menu():
    """Get a sample parsed menu for testing"""
    
    sample_text = """
    APPETIZERS
    
    Truffle Arancini - Crispy risotto balls with black truffle, parmesan cheese $14.00 (V)
    Tuna Tartare - Fresh yellowfin, avocado, citrus, sesame oil $18.00 (GF)
    Burrata Board - House made burrata, prosciutto, fig jam, grilled bread $16.00
    
    MAIN COURSES
    
    Pan Seared Salmon - Atlantic salmon, lemon herb butter, roasted vegetables $28.00 (GF)
    Wagyu Ribeye - 12oz prime cut, garlic mashed potatoes, red wine jus $45.00
    Lobster Ravioli - House made pasta, Maine lobster, saffron cream sauce $32.00
    
    DESSERTS
    
    Chocolate Lava Cake - Warm chocolate cake, vanilla ice cream $12.00 (V)
    Tiramisu - Classic Italian, espresso, mascarpone $10.00 (V)
    """
    
    menu_parser = MenuParser()
    result = menu_parser.parse_menu_text(sample_text)
    
    return {
        "sample_menu": result,
        "note": "This is a sample parsed menu to demonstrate the structure"
    }

@router.get("/test-parsing")
async def test_menu_parsing():
    """Test menu parsing with sample data"""
    sample_text = """
    APPETIZERS
    Bruschetta - Toasted bread with tomatoes, basil, and garlic $12
    Caesar Salad - Fresh romaine lettuce with parmesan cheese $14
    
    MAIN COURSES
    Grilled Salmon - Atlantic salmon with lemon butter sauce $28
    Chicken Parmesan - Breaded chicken with marinara and mozzarella $24
    """
    
    # Test both parsers
    basic_parser = MenuParser()
    basic_result = basic_parser.parse_menu_text(sample_text)
    
    enhanced_parser = EnhancedMenuParser()
    enhanced_result = enhanced_parser.parse_menu_text(sample_text, "text")
    
    return {
        "sample_text": sample_text,
        "basic_parser_result": basic_result,
        "enhanced_parser_result": enhanced_result,
        "comparison": {
            "basic_items": basic_result.get('total_items', 0),
            "enhanced_items": enhanced_result.get('total_items', 0),
            "basic_confidence": basic_result.get('parsing_info', {}).get('confidence', 0),
            "enhanced_confidence": enhanced_result.get('parsing_info', {}).get('confidence', 0)
        }
    }

@router.post("/test-word-parsing")
async def test_word_parsing(file: UploadFile = File(...)):
    """Test Word document parsing specifically"""
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Please upload a .docx file")
    
    try:
        content = await file.read()
        enhanced_parser = EnhancedMenuParser()
        result = enhanced_parser.parse_word_document(content)
        
        return {
            "filename": file.filename,
            "file_size": len(content),
            "parsing_result": result,
            "debug_info": {
                "success": result['success'],
                "total_items": result.get('total_items', 0),
                "total_sections": result.get('total_sections', 0),
                "confidence": result.get('parsing_info', {}).get('confidence', 0),
                "quality": result.get('summary', {}).get('parsing_quality', 'Unknown')
            }
        }
        
    except Exception as e:
        return {
            "filename": file.filename,
            "error": str(e),
            "success": False
        }

@router.post("/debug-text-parsing")
async def debug_text_parsing(text_data: dict):
    """Debug text parsing with detailed analysis"""
    text = text_data.get("text", "")
    
    if not text:
        raise HTTPException(status_code=400, detail="No text provided")
    
    enhanced_parser = EnhancedMenuParser()
    
    # Step-by-step analysis
    cleaned_text = enhanced_parser.clean_text(text, "text")
    sections = enhanced_parser.extract_sections_enhanced(cleaned_text)
    
    debug_info = {
        "original_text": text,
        "cleaned_text": cleaned_text,
        "sections_found": len(sections),
        "section_names": [name for name, _ in sections],
        "section_analysis": []
    }
    
    for section_name, section_text in sections:
        items = enhanced_parser.extract_menu_items_enhanced(section_text, section_name)
        debug_info["section_analysis"].append({
            "section_name": section_name,
            "section_text": section_text[:200] + "..." if len(section_text) > 200 else section_text,
            "items_found": len(items),
            "items": [
                {
                    "title": item.title,
                    "price": item.price,
                    "description": item.description,
                    "confidence": item.confidence
                }
                for item in items
            ]
        })
    
    # Full parsing result
    full_result = enhanced_parser.parse_menu_text(text, "text")
    
    return {
        "debug_info": debug_info,
        "full_result": full_result
    } 